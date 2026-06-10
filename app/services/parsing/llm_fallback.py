"""Fallback LLM: digiere cualquier formato que los parsers de codigo no reconozcan.

Extrae texto del documento (pdfplumber / python-docx / decode plano) y le pide
al modelo del rol `parser_fallback` la lista estructurada de items, validada
con Pydantic via Instructor.
"""

import io

from pydantic import BaseModel, Field

from app.core.logging import get_logger
from app.models import ParseMethod
from app.services.llm import client as llm
from app.services.parsing.base import ParsedItem, ParseResult

log = get_logger(__name__)

MAX_CHARS = 60_000


class _Item(BaseModel):
    name: str = Field(description="Nombre corto del item a cotizar")
    description: str | None = Field(default=None, description="Descripcion completa")
    quantity: int = Field(default=1, ge=1)
    preferred_brand: str | None = None
    specs: dict[str, str] = Field(
        default_factory=dict,
        description="Especificaciones tecnicas clave-valor, ej {'potencia': '750W'}",
    )


class _ItemList(BaseModel):
    items: list[_Item]


def _extract_text(filename: str, content: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
                for table in page.extract_tables():
                    for row in table:
                        parts.append(" | ".join(str(c) for c in row if c is not None))
        return "\n".join(parts)
    if name.endswith(".docx"):
        import docx

        doc = docx.Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    # texto plano / cuerpo de correo / formatos desconocidos
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="replace")


def parse_with_llm(filename: str, content: bytes, request_id: str | None = None) -> ParseResult:
    text = _extract_text(filename, content)[:MAX_CHARS]
    log.info("llm_parser_fallback", filename=filename, chars=len(text))

    result = llm.structured(
        role="parser_fallback",
        request_id=request_id,
        messages=[
            {
                "role": "system",
                "content": (
                    "Eres un extractor de solicitudes de cotizacion. Recibes el texto de un "
                    "documento o correo y devuelves la lista de items a cotizar. Extrae solo "
                    "items reales (productos/servicios solicitados), nunca inventes. Si una "
                    "cantidad no esta indicada usa 1. Pasa las especificaciones tecnicas "
                    "(potencia, medidas, material, modelo, voltaje, etc.) al dict specs."
                ),
            },
            {"role": "user", "content": f"Documento '{filename}':\n\n{text}"},
        ],
        schema=_ItemList,
    )
    items = [
        ParsedItem(
            name=i.name,
            description=i.description,
            quantity=i.quantity,
            preferred_brand=i.preferred_brand,
            specs=i.specs,
        )
        for i in result.items
    ]
    return ParseResult(items=items, method=ParseMethod.LLM_FALLBACK.value)
